"""
报销单视图函数
处理报销单的创建、编辑、查看、审核、导出等业务逻辑
"""
import io
import os
import zipfile
import openpyxl
from docx import Document
from docx.shared import Pt
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.utils import timezone
from django.db.models import Sum
from .models import Reimbursement, ReimbursementItem, Invoice, ActivityTheme
from .forms import ReimbursementForm, ReimbursementItemFormSet


@login_required
def dashboard(request):
    """仪表盘首页"""
    context = {}
    
    if request.user.is_lead:
        claims = Reimbursement.objects.filter(
            department=request.user.department
        ).exclude(status=Reimbursement.Status.DRAFT).order_by('-created_at')
        context['claims'] = claims
        context['dept_filter'] = request.user.department
    else:
        my_claims = Reimbursement.objects.filter(applicant=request.user).order_by('-created_at')
        context['my_claims'] = my_claims
        
    return render(request, 'dashboard.html', context)


@login_required
def get_activity_themes(request):
    """
    API：获取当前部门的活动主题列表
    用于前端下拉选择框的数据源
    """
    department = request.user.department or "未分配"
    
    # 查询该部门的所有活动主题
    themes = ActivityTheme.objects.filter(department=department).order_by('-created_at')
    
    data = [{
        'id': theme.id,
        'name': theme.name,
        'year': theme.activity_year,
        'month': theme.activity_month,
        'day': theme.activity_day
    } for theme in themes]
    
    return JsonResponse({'themes': data})


@login_required
def create_reimbursement(request):
    """创建新报销单"""
    if request.method == 'POST':
        form = ReimbursementForm(request.POST, department=request.user.department)
        formset = ReimbursementItemFormSet(request.POST, request.FILES)
        
        if form.is_valid() and formset.is_valid():
            reimbursement = form.save(commit=False)
            reimbursement.applicant = request.user
            reimbursement.department = request.user.department or "未分配"
            
            # 处理活动主题：查找或创建
            theme_name = form.cleaned_data['theme']
            activity_theme, created = ActivityTheme.objects.get_or_create(
                name=theme_name,
                department=reimbursement.department,
                defaults={
                    'activity_year': form.cleaned_data['activity_year'],
                    'activity_month': form.cleaned_data['activity_month'],
                    'activity_day': form.cleaned_data['activity_day'],
                }
            )
            reimbursement.activity_theme = activity_theme
            
            if 'save_draft' in request.POST:
                reimbursement.status = Reimbursement.Status.DRAFT
                messages.info(request, '报销单已暂存，您可以稍后继续编辑')
            else:
                reimbursement.status = Reimbursement.Status.SUBMITTED
                messages.success(request, '报销单提交成功！')
            
            reimbursement.save()
            formset.instance = reimbursement
            formset.save()
            
            # 处理每个物品的发票上传
            for i, item_form in enumerate(formset):
                if item_form.instance.pk:
                    files = request.FILES.getlist(f'item_{i}_files')
                    for f in files:
                        Invoice.objects.create(
                            item=item_form.instance,
                            file=f,
                            file_name=f.name
                        )
            
            reimbursement.calculate_total()
            return redirect('dashboard')
    else:
        form = ReimbursementForm(department=request.user.department)
        formset = ReimbursementItemFormSet()
    
    return render(request, 'claims/reimbursement_form.html', {
        'form': form,
        'formset': formset,
        'is_new': True
    })


@login_required
def edit_reimbursement(request, pk):
    """编辑已有报销单"""
    reimbursement = get_object_or_404(Reimbursement, pk=pk)
    
    if reimbursement.applicant != request.user:
        messages.error(request, "您没有权限编辑此报销单")
        return redirect('dashboard')
    
    if reimbursement.status not in [Reimbursement.Status.DRAFT, Reimbursement.Status.REJECTED]:
        messages.error(request, "只有草稿或已驳回状态的报销单可以编辑")
        return redirect('reimbursement_detail', pk=pk)
    
    if request.method == 'POST':
        form = ReimbursementForm(request.POST, instance=reimbursement, department=request.user.department)
        formset = ReimbursementItemFormSet(request.POST, request.FILES, instance=reimbursement)
        
        if form.is_valid() and formset.is_valid():
            reimbursement = form.save(commit=False)
            
            # 更新活动主题关联
            theme_name = form.cleaned_data['theme']
            activity_theme, created = ActivityTheme.objects.get_or_create(
                name=theme_name,
                department=reimbursement.department,
                defaults={
                    'activity_year': form.cleaned_data['activity_year'],
                    'activity_month': form.cleaned_data['activity_month'],
                    'activity_day': form.cleaned_data['activity_day'],
                }
            )
            reimbursement.activity_theme = activity_theme
            
            if 'save_draft' in request.POST:
                reimbursement.status = Reimbursement.Status.DRAFT
                messages.info(request, '报销单已暂存')
            else:
                reimbursement.status = Reimbursement.Status.SUBMITTED
                messages.success(request, '报销单已重新提交！')
            
            reimbursement.save()
            formset.save()
            
            # 处理新上传的发票
            for i, item_form in enumerate(formset):
                if item_form.instance.pk:
                    files = request.FILES.getlist(f'item_{i}_files')
                    for f in files:
                        Invoice.objects.create(
                            item=item_form.instance,
                            file=f,
                            file_name=f.name
                        )
            
            reimbursement.calculate_total()
            return redirect('dashboard')
    else:
        form = ReimbursementForm(instance=reimbursement, department=request.user.department)
        formset = ReimbursementItemFormSet(instance=reimbursement)
    
    return render(request, 'claims/reimbursement_form.html', {
        'form': form,
        'formset': formset,
        'reimbursement': reimbursement,
        'is_new': False
    })


@login_required
def reimbursement_detail(request, pk):
    """查看报销单详情"""
    reimbursement = get_object_or_404(Reimbursement, pk=pk)
    
    if not request.user.is_lead and reimbursement.applicant != request.user:
        messages.error(request, "您没有权限查看此报销单")
        return redirect('dashboard')

    return render(request, 'claims/reimbursement_detail.html', {
        'reimbursement': reimbursement
    })


@login_required
def review_reimbursement(request, pk):
    """审核报销单（仅负责人可用）"""
    if not request.user.is_lead:
        return redirect('dashboard')
        
    reimbursement = get_object_or_404(Reimbursement, pk=pk)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        note = request.POST.get('note', '')
        
        reimbursement.reviewer_note = note
        
        if action == 'approve':
            reimbursement.status = Reimbursement.Status.PACKED
            messages.success(request, '已标记为已打包')
        elif action == 'reject':
            reimbursement.status = Reimbursement.Status.REJECTED
            messages.warning(request, '已驳回该申请')
            
        reimbursement.save()
        return redirect('dashboard')
        
    return redirect('reimbursement_detail', pk=pk)


@login_required
def export_claims(request):
    """
    导出报销数据（仅负责人可用）
    按活动主题分组，生成：
    1. Excel汇总表
    2. 每个主题一个Word说明文档
    3. 按主题和物品分层存放的票据文件夹
    """
    if not request.user.is_lead:
        messages.error(request, "无权限执行此操作")
        return redirect('dashboard')

    department = request.user.department
    
    # 获取本部门所有已提交/已打包的报销单
    claims = Reimbursement.objects.filter(
        department=department
    ).exclude(status=Reimbursement.Status.DRAFT).exclude(status=Reimbursement.Status.REJECTED)

    if not claims.exists():
        messages.warning(request, "没有可导出的报销单")
        return redirect('dashboard')

    # 按活动主题分组
    theme_groups = {}
    for claim in claims:
        theme_name = claim.theme
        if theme_name not in theme_groups:
            theme_groups[theme_name] = {
                'claims': [],
                'leaders': set(),
                'locations': set(),
                'activity_date': claim.activity_date_display,
                'total_amount': 0
            }
        theme_groups[theme_name]['claims'].append(claim)
        theme_groups[theme_name]['leaders'].add(claim.activity_leader)
        if claim.activity_location:
            theme_groups[theme_name]['locations'].add(claim.activity_location)
        theme_groups[theme_name]['total_amount'] += float(claim.total_amount)

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        
        # ========== 1. 生成Excel汇总表 ==========
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "报销汇总"
        ws.append(['序号', '活动主题', '申请人', '学号', '物品名称', '数量', '单位', '单价', '金额', '活动时间', '活动地点'])
        
        row_num = 1
        for theme_name, group in theme_groups.items():
            for claim in group['claims']:
                for item in claim.items.all():
                    ws.append([
                        row_num,
                        theme_name,
                        claim.applicant.first_name or claim.applicant.username,
                        claim.applicant.student_id or '',
                        item.name,
                        item.quantity,
                        item.unit,
                        float(item.price),
                        float(item.amount),
                        claim.activity_date_display,
                        claim.activity_location
                    ])
                    row_num += 1
        
        excel_buffer = io.BytesIO()
        wb.save(excel_buffer)
        excel_name = f"01_环境学院xx年xx学期第xx次报账_{department}.xlsx"
        zip_file.writestr(excel_name, excel_buffer.getvalue())
        
        # ========== 2. 为每个主题生成Word说明文档 ==========
        theme_index = 0
        for theme_name, group in theme_groups.items():
            theme_index += 1
            
            doc = Document()
            doc.add_heading('学生活动经费使用情况说明', 0)
            
            # 基本信息
            doc.add_paragraph(f"活动主题：{theme_name}")
            doc.add_paragraph(f"参与人员：{'、'.join(group['leaders'])}")
            doc.add_paragraph(f"经办人姓名：{request.user.first_name or request.user.username}")
            doc.add_paragraph(f"经办人联系方式：")
            doc.add_paragraph(f"活动时间：{group['activity_date']}")
            doc.add_paragraph(f"活动地点：{'、'.join(group['locations'])}")
            
            # 活动内容（合并所有描述）
            doc.add_paragraph("活动主要内容：")
            for claim in group['claims']:
                if claim.description:
                    doc.add_paragraph(claim.description)
            
            # 报销内容及金额
            doc.add_paragraph("")
            doc.add_paragraph("报销内容及金额：")
            
            item_num = 0
            for claim in group['claims']:
                item_num += 1
                doc.add_paragraph(f"{item_num}. {claim.applicant.first_name or claim.applicant.username}：")
                
                for item in claim.items.all():
                    line = f"    {item.name} {item.quantity}{item.unit} —— ¥{item.amount}"
                    doc.add_paragraph(line)
            
            doc.add_paragraph(f"合计：¥{group['total_amount']:.2f}")
            
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_name = f"02_学生活动经费使用情况说明_{theme_name}.docx"
            zip_file.writestr(doc_name, doc_buffer.getvalue())
            
            # ========== 3. 打包票据文件 ==========
            # 票据/1主题名/1.1_物品名/发票.pdf
            item_global_index = 0
            for claim in group['claims']:
                for item in claim.items.all():
                    item_global_index += 1
                    
                    # 文件夹路径
                    folder_path = f"票据/{theme_index}{theme_name}/{theme_index}.{item_global_index}_{item.name}"
                    
                    for invoice in item.invoices.all():
                        try:
                            file_path = invoice.file.path
                            file_name = invoice.file_name or invoice.file.name.split('/')[-1]
                            arcname = f"{folder_path}/{file_name}"
                            zip_file.write(file_path, arcname=arcname)
                        except Exception as e:
                            print(f"Error packing file: {e}")

    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/zip')
    filename = f"报销导出_{department}_{timezone.now().strftime('%Y%m%d%H%M')}.zip"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def delete_invoice(request, invoice_id):
    """删除单个发票凭证"""
    invoice = get_object_or_404(Invoice, pk=invoice_id)
    
    # 检查权限
    if invoice.item.reimbursement.applicant != request.user:
        return JsonResponse({'error': '无权限'}, status=403)
    
    # 删除文件
    if invoice.file:
        try:
            os.remove(invoice.file.path)
        except:
            pass
    
    invoice.delete()
    return JsonResponse({'success': True})
